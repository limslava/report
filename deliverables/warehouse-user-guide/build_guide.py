from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
ASSETS = ROOT / "assets"
OUTPUT = ROOT / "Инструкция_пользователя_Склад_ТС.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(35, 45, 58)
MUTED = RGBColor(98, 108, 120)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_BLUE = "F4F7FB"
WHITE = "FFFFFF"
RED = RGBColor(155, 28, 28)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if table.rows:
        header_pr = table.rows[0]._tr.get_or_add_trPr()
        if header_pr.find(qn("w:tblHeader")) is None:
            header_pr.append(OxmlElement("w:tblHeader"))


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Страница ")
    set_run(run, size=9, color=MUTED)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def set_run(run, size=None, color=None, bold=None, italic=None, font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), font)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    title.font.size = Pt(30)
    title.font.bold = True
    title.font.color.rgb = DARK_BLUE
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)

    subtitle = doc.styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    subtitle._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    subtitle.font.size = Pt(15)
    subtitle.font.color.rgb = MUTED
    subtitle.paragraph_format.space_after = Pt(8)

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_bullet(doc, text):
    return doc.add_paragraph(text, style="List Bullet")


def add_steps(doc, items):
    numbering = doc.part.numbering_part.element
    base_num_id = int(doc.styles["List Number"]._element.pPr.numPr.numId.val)
    base_num = next(
        num for num in numbering.findall(qn("w:num"))
        if int(num.get(qn("w:numId"))) == base_num_id
    )
    abstract_num_id = base_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    existing_ids = [int(num.get(qn("w:numId"))) for num in numbering.findall(qn("w:num"))]
    num_id = max(existing_ids) + 1

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_num_id)
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)

    for text in items:
        p = doc.add_paragraph(style="List Number")
        num_pr = p._p.get_or_add_pPr().get_or_add_numPr()
        num_pr.get_or_add_ilvl().val = 0
        num_pr.get_or_add_numId().val = num_id
        p.add_run(text)


def add_note(doc, title, text, kind="info"):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "FFF4E5" if kind == "warning" else PALE_BLUE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{title}. ")
    set_run(r, bold=True, color=RED if kind == "warning" else DARK_BLUE)
    r = p.add_run(text)
    set_run(r, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_figure(doc, filename, caption, width=6.2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run()
    picture = r.add_picture(str(ASSETS / filename), width=Inches(width))
    picture._inline.docPr.set("descr", caption)
    picture._inline.docPr.set("title", caption)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(0)
    cap.paragraph_format.space_after = Pt(9)
    cap.paragraph_format.keep_with_next = False
    r = cap.add_run(caption)
    set_run(r, size=9, color=MUTED, italic=True)


def add_role_table(doc):
    rows = [
        (
            "Администратор",
            "Полный доступ к модулю: реестр, приёмка, выдача, фото, услуги, клиенты, тарифы, начисления, акты, закрытие периода и корректировка дат.",
            "Корректировка дат выполняется только через специальное действие, с обязательной причиной и записью в аудит.",
        ),
        (
            "Руководитель склада",
            "Все складские операции; управление клиентами и тарифами; контроль услуг и начислений; закрытие периода; контролируемая корректировка даты и времени.",
            "Нельзя менять дату приёмки/выдачи обычным редактированием карточки. Закрытый период защищён от изменений.",
        ),
        (
            "Кладовщик",
            "Приёмка и выдача ТС, фотофиксация, поиск техники на стоянке, фиксация дополнительных услуг и их количества.",
            "Не управляет клиентами и тарифами, не видит финансовый реестр и акты, не корректирует дату/время операций.",
        ),
        (
            "Финансист",
            "Тарифы, дополнительные услуги, расчёт начислений, Excel-реестр, PDF-акт и закрытие периода.",
            "Не принимает и не выдаёт ТС, не загружает складские фотографии, не корректирует время приёмки/выдачи.",
        ),
        (
            "Директор / генеральный директор",
            "Просмотр реестра и финансовых начислений для контроля.",
            "Режим просмотра: без приёмки, выдачи, изменения услуг, тарифов, дат и закрытия периода.",
        ),
        (
            "Представитель контрагента",
            "Просмотр только своей техники, связанных услуг, начислений и актов.",
            "Не видит технику других контрагентов; не создаёт и не редактирует ТС, услуги, тарифы и акты.",
        ),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_geometry(table, [1720, 4240, 3400])
    headers = ("Роль", "Возможности", "Ограничения")
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text)
        set_run(r, size=9.5, bold=True, color=DARK_BLUE)
    for role, abilities, limits in rows:
        cells = table.add_row().cells
        for idx, text in enumerate((role, abilities, limits)):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(text)
            set_run(r, size=8.5, bold=(idx == 0), color=INK)
    set_table_geometry(table, [1720, 4240, 3400])


def add_access_table(doc):
    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    widths = [1800, 1260, 1260, 1260, 1260, 1260, 1260]
    set_table_geometry(table, widths)
    headers = ("Операция", "Админ", "Рук. склада", "Кладовщик", "Финансист", "Директор", "Контрагент")
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text)
        set_run(r, size=7.5, bold=True, color=DARK_BLUE)
    data = [
        ("Приёмка / выдача", "Да", "Да", "Да", "Нет", "Нет", "Нет"),
        ("Фото ТС", "Да", "Да", "Да", "Нет", "Нет", "Просмотр своих*"),
        ("Фиксация услуг", "Да", "Да", "Да", "Да", "Нет", "Нет"),
        ("Клиенты склада", "Да", "Да", "Нет", "Нет", "Нет", "Нет"),
        ("Тарифы", "Да", "Да", "Нет", "Да", "Нет", "Нет"),
        ("Начисления / акты", "Да", "Да", "Нет", "Да", "Просмотр", "Только свои"),
        ("Закрытие периода", "Да", "Да", "Нет", "Да", "Нет", "Нет"),
        ("Корректировка дат", "Да", "Да", "Нет", "Нет", "Нет", "Нет"),
    ]
    for row_data in data:
        cells = table.add_row().cells
        for idx, text in enumerate(row_data):
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            r = p.add_run(text)
            set_run(r, size=7.5, bold=(idx == 0), color=INK)
    set_table_geometry(table, widths)


def add_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Склад ТС  |  Инструкция пользователя")
    set_run(r, size=9, color=MUTED)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "D9E2F3")
    pBdr.append(bottom)
    pPr.append(pBdr)
    add_page_number(section.footer.paragraphs[0])


def build_document():
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    add_header_footer(section)

    # Cover
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(112)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("РУКОВОДСТВО ПОЛЬЗОВАТЕЛЯ")
    set_run(r, size=11, bold=True, color=BLUE)
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Склад транспортных средств")
    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Приёмка, хранение, дополнительные услуги, выдача и взаиморасчёты")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    r = p.add_run("Версия инструкции: 23 июня 2026 года")
    set_run(r, size=10, color=MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Модуль находится в разработке; состав экранов может уточняться до выпуска.")
    set_run(r, size=9.5, color=MUTED, italic=True)
    doc.add_page_break()

    add_heading(doc, "1. Назначение инструкции", 1)
    doc.add_paragraph(
        "Инструкция описывает текущую работу модуля «Склад ТС» и предназначена для "
        "администратора, руководителя склада, кладовщика, финансиста, руководителей компании "
        "и представителей контрагентов."
    )
    add_note(
        doc,
        "Главный принцип",
        "пользователь видит только те разделы и действия, которые разрешены его ролью. "
        "Отсутствие кнопки — штатное ограничение доступа, а не ошибка системы.",
    )

    add_heading(doc, "2. Основные правила учёта", 1)
    for text in (
        "Хранение рассчитывается по календарным датам; день приёмки и день выдачи включаются.",
        "Тариф хранения и стоимость услуг зависят от типа ТС: легковой или грузовой.",
        "Дата и время фактической приёмки и выдачи проставляются системой автоматически.",
        "Одна и та же дополнительная услуга может выполняться несколько раз.",
        "Для дозаправки кладовщик указывает фактическое количество литров; цена применяется за литр.",
        "Цена услуги фиксируется в момент выполнения и используется в начислениях.",
        "При выдаче обязательна фотофиксация. Фотографии хранятся, пока ТС находится на стоянке, "
        "и удаляются после выдачи; техническая резервная копия создаётся автоматически.",
        "После закрытия расчётного периода начисления и даты защищены от изменений.",
    ):
        add_bullet(doc, text)

    add_heading(doc, "3. Роли, права и ограничения", 1)
    add_role_table(doc)
    doc.add_paragraph()
    add_heading(doc, "Краткая матрица доступа", 2)
    add_access_table(doc)
    p = doc.add_paragraph()
    r = p.add_run("* ")
    set_run(r, size=9, color=MUTED)
    r = p.add_run("Фактическая доступность фотографий контрагенту определяется интерфейсом его карточки и политикой хранения.")
    set_run(r, size=9, color=MUTED)

    doc.add_page_break()
    add_heading(doc, "4. Вход и навигация", 1)
    add_steps(doc, [
        "Откройте адрес системы, выданный администратором.",
        "Введите логин и пароль своей учётной записи.",
        "В боковом меню выберите «Склад ТС» или «Рабочая станция» — название зависит от роли.",
        "Для завершения работы используйте пункт «Выход». Не передавайте свою учётную запись другим сотрудникам.",
    ])
    add_figure(doc, "01-registry-admin.jpg", "Рисунок 1. Реестр ТС в режиме администратора.")
    add_note(
        doc,
        "Поиск",
        "реестр позволяет искать ТС по складскому номеру, VIN, государственному номеру и марке, "
        "а также фильтровать список по статусу и типу ТС.",
    )

    add_heading(doc, "5. Работа кладовщика", 1)
    doc.add_paragraph(
        "Кладовщик работает преимущественно через «Рабочую станцию кладовщика». "
        "На одном экране доступны приёмка, выдача, поиск техники на стоянке и быстрый переход к услугам."
    )
    add_figure(doc, "05-operations.jpg", "Рисунок 2. Рабочая станция кладовщика.")

    add_heading(doc, "5.1. Приёмка ТС", 2)
    add_steps(doc, [
        "Нажмите «Принять ТС».",
        "На шаге «Основа» выберите активного клиента склада, укажите номер и дату заявки.",
        "На шаге «ТС» заполните тип, марку, модель и идентификаторы ТС. Проверьте VIN, номер шасси и госномер.",
        "На шаге «Осмотр» зафиксируйте уровень топлива и комментарий о состоянии.",
        "На шаге «Фото» выполните фотофиксацию состояния ТС. Снимки должны позволять различить повреждения.",
        "На шаге «Проверка» сверьте данные и подтвердите приёмку.",
    ])
    add_figure(doc, "06-reception-wizard.jpg", "Рисунок 3. Пошаговая форма приёмки ТС.")
    add_note(
        doc,
        "Дата и время",
        "фактический момент приёмки задаёт сервер. Кладовщик не может выбрать прошедшее или будущее время вручную.",
        kind="warning",
    )
    add_note(
        doc,
        "Черновик",
        "форма приёмки автоматически сохраняет черновик на текущем устройстве. "
        "Перед началом новой приёмки убедитесь, что предыдущий черновик завершён или очищен.",
    )

    add_heading(doc, "5.2. Дополнительные услуги", 2)
    add_steps(doc, [
        "Найдите ТС на стоянке по складскому номеру, VIN, госномеру, марке или контрагенту.",
        "Выберите ТС и нажмите «Открыть услуги» либо используйте кнопку услуг в реестре.",
        "Отметьте выполненную услугу, укажите количество и комментарий, если это требуется.",
        "Для дозаправки укажите фактическое количество литров.",
        "Сохраните выполнение. Повторную операцию оформляйте новой записью, не изменяя предыдущую.",
    ])
    add_note(
        doc,
        "Фиксация стоимости",
        "система сохраняет дату, исполнителя, количество, действующую цену и итоговую сумму. "
        "Это позволяет восстановить расчёт даже после изменения тарифа.",
    )

    add_heading(doc, "5.3. Выдача ТС", 2)
    add_steps(doc, [
        "Нажмите «Выдать ТС».",
        "Найдите и выберите ТС со статусом «На стоянке».",
        "Проверьте карточку и список оказанных услуг.",
        "Добавьте обязательные фотографии состояния ТС при выдаче.",
        "На шаге подтверждения завершите выдачу. Дата и время будут проставлены автоматически.",
    ])
    add_figure(doc, "07-issue-wizard.jpg", "Рисунок 4. Пошаговая форма выдачи ТС.")
    add_note(
        doc,
        "Важно",
        "после подтверждения выдачи фотографии удаляются согласно принятому сроку хранения. "
        "Не завершайте операцию, пока проверка и фотофиксация не закончены.",
        kind="warning",
    )

    add_heading(doc, "6. Работа руководителя склада", 1)
    doc.add_paragraph(
        "Руководитель склада сочетает операционные и контрольные функции: может выполнять складские операции, "
        "управлять клиентами и тарифами, контролировать начисления и исправлять дату/время через специальную процедуру."
    )
    add_heading(doc, "6.1. Клиенты склада", 2)
    add_steps(doc, [
        "Откройте вкладку «Клиенты склада».",
        "Добавьте организацию, которая действительно передаёт ТС на хранение.",
        "Укажите реквизиты договора хранения и дату начала работы.",
        "Оставляйте клиента активным только на период фактического обслуживания.",
    ])
    add_figure(doc, "02-clients.jpg", "Рисунок 5. Отдельный справочник клиентов склада.")
    add_note(
        doc,
        "Почему отдельный справочник",
        "не каждый контрагент из договоров является клиентом склада. "
        "В форме приёмки показываются только активные клиенты склада.",
    )

    add_heading(doc, "6.2. Контролируемая корректировка даты и времени", 2)
    add_steps(doc, [
        "В реестре откройте действие «Скорректировать дату и время».",
        "Укажите правильное значение и обязательную причину корректировки.",
        "Подтвердите изменение.",
    ])
    add_note(
        doc,
        "Аудит",
        "система сохраняет старое и новое значение, причину и пользователя, выполнившего корректировку. "
        "Корректировка запрещена для закрытого периода и не может устанавливать время в будущем.",
        kind="warning",
    )

    add_heading(doc, "7. Работа финансиста", 1)
    add_heading(doc, "7.1. Тарифы", 2)
    add_steps(doc, [
        "Откройте вкладку «Услуги и тарифы».",
        "Выберите услугу и нажмите «Настроить».",
        "Задайте цену отдельно для легкового и грузового ТС и дату начала действия.",
        "Для топлива указывайте цену за литр; количество вводит кладовщик при выполнении.",
    ])
    add_figure(doc, "03-tariffs.jpg", "Рисунок 6. Справочник услуг и тарифов.")
    add_note(
        doc,
        "История тарифов",
        "новая цена действует с указанной даты. Расчёт хранения учитывает тариф, действовавший в каждый календарный день.",
    )

    add_heading(doc, "7.2. Начисления, Excel и акт", 2)
    add_steps(doc, [
        "Откройте вкладку «Начисления и акты».",
        "Выберите период, контрагента и при необходимости тип ТС.",
        "Нажмите «Рассчитать» и проверьте количество ТС, суток хранения, услуги и итог.",
        "Скачайте Excel-реестр. Выгрузка может формироваться по всем контрагентам.",
        "Для PDF-акта выберите одного контрагента.",
        "После проверки закройте период. Закрытие фиксирует расчёт и запрещает последующие изменения.",
    ])
    add_figure(doc, "04-billing.jpg", "Рисунок 7. Предварительный расчёт начислений и формирование документов.")
    add_note(
        doc,
        "Формула",
        "итого по ТС = хранение за включённые календарные сутки + сумма всех дополнительных услуг.",
    )
    add_note(
        doc,
        "Перед закрытием",
        "проверьте выбранного контрагента, период, тарифы, даты приёмки/выдачи и перечень услуг. "
        "Закрытый период предназначен для окончательных взаиморасчётов.",
        kind="warning",
    )

    add_heading(doc, "8. Руководители компании", 1)
    doc.add_paragraph(
        "Директор и генеральный директор используют модуль для контроля. Им доступен просмотр реестра и начислений, "
        "но недоступны операционные и финансовые изменения."
    )
    for text in (
        "можно искать ТС и анализировать статусы;",
        "можно просматривать расчёт хранения и дополнительных услуг;",
        "нельзя принимать или выдавать ТС;",
        "нельзя менять услуги, тарифы, даты и закрывать период.",
    ):
        add_bullet(doc, text)

    add_heading(doc, "9. Представитель контрагента", 1)
    doc.add_paragraph(
        "Учётная запись представителя привязывается администратором к одному клиенту склада. "
        "Ограничение применяется на сервере: пользователь получает только данные своего контрагента."
    )
    add_steps(doc, [
        "Войдите под выданной учётной записью.",
        "Откройте «Склад ТС» и просмотрите собственную технику.",
        "Откройте начисления и акты; выбор другого контрагента недоступен.",
    ])
    for text in (
        "кнопка приёмки отсутствует;",
        "редактирование карточки, услуг и дат недоступно;",
        "техника и финансовые данные других клиентов не отображаются;",
        "при попытке прямого обращения к запрещённой операции сервер отклоняет запрос.",
    ):
        add_bullet(doc, text)

    add_heading(doc, "10. Фотофиксация и хранение файлов", 1)
    for text in (
        "Снимайте общий вид, четыре стороны, идентификаторы, салон и все обнаруженные повреждения.",
        "Не используйте чрезмерное сжатие: мелкие царапины и дефекты должны оставаться различимыми.",
        "При выдаче создаётся отдельный комплект фотографий.",
        "Файлы сохраняются в основном хранилище и автоматически дублируются в резервное.",
        "После выдачи ТС фотографии удаляются в соответствии с принятой политикой хранения.",
    ):
        add_bullet(doc, text)
    add_note(
        doc,
        "Организационная мера",
        "резервное хранилище должно находиться на отдельном диске или сервере и регулярно проверяться ответственным сотрудником.",
        kind="warning",
    )

    add_heading(doc, "11. Типовые ситуации", 1)
    situations = [
        ("ТС не находится при выдаче", "Проверьте статус: выдавать можно только ТС «На стоянке». Очистите поиск и попробуйте складской номер или VIN."),
        ("Контрагент отсутствует в приёмке", "Попросите руководителя склада проверить, добавлен ли клиент в отдельный справочник и активен ли он."),
        ("Нельзя сформировать PDF-акт", "Выберите одного контрагента. По всем контрагентам доступна Excel-выгрузка."),
        ("Нельзя изменить дату", "Кладовщик не имеет такого права. Руководитель склада выполняет корректировку с указанием причины."),
        ("Нельзя исправить закрытый период", "Это штатная защита. Необходимо организационное решение ответственных лиц, а не обычное редактирование."),
        ("Кнопка действия не отображается", "Проверьте роль пользователя. Интерфейс скрывает запрещённые операции."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [3000, 6360])
    for idx, text in enumerate(("Ситуация", "Что делать")):
        set_cell_shading(table.rows[0].cells[idx], LIGHT_BLUE)
        r = table.rows[0].cells[idx].paragraphs[0].add_run(text)
        set_run(r, size=9.5, bold=True, color=DARK_BLUE)
    for problem, solution in situations:
        cells = table.add_row().cells
        for idx, text in enumerate((problem, solution)):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            r = p.add_run(text)
            set_run(r, size=9, bold=(idx == 0), color=INK)
    set_table_geometry(table, [3000, 6360])

    add_heading(doc, "12. Ответственность пользователей", 1)
    for text in (
        "Кладовщик отвечает за точность карточки, фактическое количество услуг и качество фотофиксации.",
        "Руководитель склада отвечает за справочник клиентов, обоснованность корректировок и операционный контроль.",
        "Финансист отвечает за тарифы, проверку начислений, документы и своевременное закрытие периода.",
        "Администратор отвечает за учётные записи, роли и правильную привязку представителя к контрагенту.",
        "Каждый пользователь отвечает за сохранность своей учётной записи и действия, записанные в аудит.",
    ):
        add_bullet(doc, text)

    add_note(
        doc,
        "Статус документа",
        "инструкция отражает функциональность модуля на 23 июня 2026 года в отдельной ветке разработки. "
        "Перед промышленным запуском документ следует актуализировать после окончательной приёмки.",
    )

    core = doc.core_properties
    core.title = "Инструкция пользователя — Склад транспортных средств"
    core.subject = "Роли, права, ограничения и пользовательские сценарии"
    core.author = "Проект Report"
    core.keywords = "склад ТС, приёмка, выдача, услуги, тарифы, начисления, роли"
    core.comments = "Сформировано по текущей версии модуля склада."

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document())
