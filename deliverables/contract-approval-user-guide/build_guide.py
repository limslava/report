from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
ASSETS = ROOT / "assets"
OUTPUT = ROOT / "Инструкция_пользователя_БП_договоров.docx"
DOCS_OUTPUT = ROOT.parents[1] / "docs" / "user-guides" / "contract_approval_user_guide.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(35, 45, 58)
MUTED = RGBColor(98, 108, 120)
RED = RGBColor(155, 28, 28)
GREEN = RGBColor(38, 116, 51)
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F4F7FB"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"


def set_run(run, size=None, color=None, bold=None, italic=None, font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), font)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if table.rows:
        header_pr = table.rows[0]._tr.get_or_add_trPr()
        if header_pr.find(qn("w:tblHeader")) is None:
            header_pr.append(OxmlElement("w:tblHeader"))


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.18

    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    title.font.size = Pt(26)
    title.font.bold = True
    title.font.color.rgb = DARK_BLUE
    title.paragraph_format.space_after = Pt(6)

    subtitle = doc.styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    subtitle._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    subtitle.font.size = Pt(13)
    subtitle.font.color.rgb = MUTED
    subtitle.paragraph_format.space_after = Pt(8)

    for name, size, color, before, after in (
        ("Heading 1", 15, BLUE, 14, 7),
        ("Heading 2", 12, BLUE, 10, 5),
        ("Heading 3", 11, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(10)
        style.paragraph_format.left_indent = Inches(0.35)
        style.paragraph_format.first_line_indent = Inches(-0.18)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.18


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Страница ")
    set_run(run, size=9, color=MUTED)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def add_title(doc):
    p = doc.add_paragraph(style="Title")
    p.add_run("Инструкция пользователя").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2 = doc.add_paragraph("Модуль «Согласование договоров»", style="Subtitle")
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_note(
        doc,
        "Для кого инструкция",
        "Документ описывает рабочие действия инициатора договора, руководителя службы безопасности, юриста, финансового директора, главного бухгалтера и офис-менеджера. Технические настройки и управление пользователями здесь не рассматриваются.",
        fill=PALE_BLUE,
    )


def add_note(doc, title, text, fill=PALE_BLUE, color=INK):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [13440])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    run = p.add_run(title)
    set_run(run, size=10, color=color, bold=True)
    p2 = cell.add_paragraph(text)
    p2.paragraph_format.space_after = Pt(0)
    for r in p2.runs:
        set_run(r, size=10, color=color)
    doc.add_paragraph()


def add_step_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [650, 3150, 9640])
    headers = ["Шаг", "Кто выполняет", "Что происходит"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, LIGHT_BLUE)
        run = cell.paragraphs[0].add_run(h)
        set_run(run, size=9, bold=True, color=DARK_BLUE)
    for idx, (owner, action) in enumerate(rows, start=1):
        cells = table.add_row().cells
        cells[0].text = str(idx)
        cells[1].text = owner
        cells[2].text = action
        for cell in cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    set_run(r, size=9)
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_steps(doc, items):
    for idx, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.35)
        p.paragraph_format.first_line_indent = Inches(-0.22)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f"{idx}.  {item}")
        set_run(run, size=10, color=INK)


def add_image(doc, filename, caption, width=9.7):
    path = ASSETS / filename
    if not path.exists():
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    for run in cap.runs:
        set_run(run, size=9, color=MUTED, italic=True)


def add_section_break(doc):
    doc.add_page_break()


def build():
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)

    configure_styles(doc)
    header_p = section.header.paragraphs[0]
    header_p.text = "Инструкция пользователя: согласование договоров"
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run(header_p.runs[0], size=9, color=MUTED)
    add_page_number(section.footer.paragraphs[0])

    add_title(doc)

    doc.add_heading("1. Общая схема процесса", level=1)
    add_step_table(doc, [
        ("Инициатор", "Создаёт договор или дополнительное соглашение, заполняет карточку, прикладывает файлы и отправляет на согласование."),
        ("Руководитель службы безопасности", "Проверяет договор и ставит визу: согласован, согласован с замечаниями или не согласован."),
        ("Юрист, финансовый директор, главный бухгалтер", "Для расходных договоров и доходных договоров с ПСР согласуют договор параллельно. Каждый участник ставит свою визу независимо от остальных."),
        ("Офис-менеджер", "После сбора обязательных виз распечатывает пакет, передаёт экземпляр генеральному директору на подпись, прикладывает подписанный скан и фиксирует итоговое решение."),
        ("Участники процесса", "Используют обсуждение в карточке договора для уточнений и файлов, не перезапуская согласование."),
    ])
    add_note(
        doc,
        "Главное правило",
        "После того как обязательный круг участников поставил визы, договор переходит офис-менеджеру независимо от того, были визы положительными, с замечаниями или отрицательными. Финальный результат фиксируется после физической подписи или отказа от подписания.",
    )

    doc.add_heading("2. Рабочий экран инициатора", level=1)
    doc.add_paragraph("На вкладке «Мои договоры» видны договоры и дополнительные соглашения, с которыми работает инициатор. Строка открывается двойным кликом. Кнопки «Добавить» и «Импорт» раскрывают меню действий.")
    add_image(doc, "01-my-contracts.png", "Вкладка «Мои договоры»: рабочий список договоров инициатора.")
    add_image(doc, "02-add-menu.png", "Меню «Добавить»: создание основного договора или дополнительного соглашения.", width=9.3)

    add_section_break(doc)
    doc.add_heading("3. Создание договора", level=1)
    doc.add_paragraph("Новый договор создаётся через «Добавить» → «Договор». Первый экран нужен для определения контрагента и типа документа.")
    add_image(doc, "03-new-contract-step-inn.png", "Первый шаг мастера: ИНН, контрагент, вид документа и тип договора.", width=6.0)
    doc.add_heading("Порядок действий инициатора", level=2)
    add_steps(doc, [
        "Введите ИНН контрагента. Если данные из ФНС подтянулись не полностью, заполните доступные поля вручную.",
        "Проверьте наименование и краткое наименование контрагента.",
        "Выберите тип договора: расходный или доходный. Для доходного договора выберите подтип: с ПСР или без ПСР.",
        "Нажмите «Проверить». Если система нашла похожие договоры, откройте их и убедитесь, что новый договор действительно нужен.",
        "Заполните номер, дату, предмет договора и способ подписания. Для доходных договоров номер присваивается автоматически при отправке.",
    ])

    doc.add_heading("Доходный договор и ПСР", level=2)
    doc.add_paragraph("Для доходного договора система формирует файл договора по проформе. Если договор с ПСР, на шаге файлов приложите ПСР клиента и другие сопроводительные документы.")
    add_image(doc, "13-income-requisites.png", "Шаг реквизитов доходного договора: данные для автогенерации документа.", width=6.0)
    add_image(doc, "14-income-files-step.png", "Шаг файлов для доходного договора с ПСР: договор сформируется автоматически, ПСР прикладывается отдельно.", width=6.0)

    add_section_break(doc)
    doc.add_heading("4. Дополнительные соглашения и импорт", level=1)
    add_bullets(doc, [
        "Дополнительное соглашение создаётся через «Добавить» → «Доп. соглашение».",
        "При создании дополнительного соглашения выберите основной договор, к которому оно относится. Список ограничивается договорами того же контрагента и подходящего типа.",
        "Если договор или дополнительное соглашение уже подписаны ранее, используйте «Импорт»: заполните карточку и обязательно приложите подписанный файл.",
        "Импортированный документ не проходит согласование; в истории решений будет указано, что документ внесён как ранее подписанный.",
    ])
    add_step_table(doc, [
        ("Инициатор", "Для дополнительного соглашения выбирает основной договор, проверяет контрагента и прикладывает файл соглашения."),
        ("Инициатор", "Для импортированного подписанного документа заполняет карточку, прикладывает подписанный файл и сохраняет запись в реестр."),
        ("Участники процесса", "По импортированным документам не ставят визы: такие записи нужны для архива и поиска подписанных документов."),
    ])

    add_section_break(doc)
    doc.add_heading("5. Карточка договора", level=1)
    doc.add_paragraph("Карточка открывается двойным кликом по строке. В ней видны реквизиты, файлы, ход согласования, история решений и обсуждение.")
    add_image(doc, "04-contract-card-overview.png", "Карточка подписанного договора: лист согласования, файлы и ход процесса.", width=8.6)
    add_bullets(doc, [
        "Кнопка «История решений» показывает журнал виз и изменений.",
        "Файлы договора открываются из блока «Документы договора».",
        "Ход согласования показывает, кто уже поставил визу и какие комментарии оставлены.",
    ])

    add_section_break(doc)
    doc.add_heading("6. Обсуждение договора", level=1)
    doc.add_paragraph("Обсуждение используется для рабочих уточнений без возврата договора на новый круг: запросить приложение, уточнить реквизиты, приложить поясняющий файл.")
    add_image(doc, "05-contract-discussion.png", "Обсуждение в карточке: сообщения, вложения и поле ввода.", width=8.7)
    add_bullets(doc, [
        "Введите текст сообщения в поле «Сообщение».",
        "Для упоминания участника начните ввод с символа @ и выберите человека из списка.",
        "К сообщению можно приложить файлы: PDF, DOC, DOCX, PNG, JPG/JPEG.",
        "После финального статуса договора обсуждение остаётся доступным только для чтения.",
    ])

    doc.add_heading("7. Работа руководителя службы безопасности", level=1)
    doc.add_paragraph("На вкладке «Согласование договоров» руководитель службы безопасности видит договоры, ожидающие проверки. Договор открывается двойным кликом.")
    add_image(doc, "06-security-inbox.png", "Список договоров на проверке руководителя службы безопасности.")
    add_image(doc, "07-security-decision-card.png", "Карточка задачи руководителя службы безопасности: выбор решения, комментарий, прикрепление файла.", width=8.7)
    add_bullets(doc, [
        "Выберите решение: «Согласован», «Согласован с замечаниями» или «Не согласован».",
        "При решении с замечаниями заполните комментарий.",
        "Если нужно приложить подтверждающий документ, используйте «Прикрепить файл» в строке своей визы.",
    ])

    add_section_break(doc)
    doc.add_heading("8. Работа юриста, финансового директора и главного бухгалтера", level=1)
    doc.add_paragraph("Для этих участников экран одинаковый: список договоров на согласовании и карточка с выбором решения. Для расходных договоров и доходных договоров с ПСР задачи идут параллельно.")
    add_image(doc, "08-approval-inbox.png", "Список договоров у согласующего участника.")
    add_image(doc, "09-approval-decision-card.png", "Карточка согласующего: решение, комментарий и прикрепление файла.", width=8.7)
    add_bullets(doc, [
        "Откройте договор двойным кликом.",
        "Просмотрите файлы договора и ход согласования.",
        "Выберите решение. Если выбираете «Согласован с замечаниями», комментарий обязателен.",
        "Нажмите «Сохранить решение».",
    ])

    add_section_break(doc)
    doc.add_heading("9. Работа офис-менеджера", level=1)
    doc.add_paragraph("Когда обязательные визы собраны, договор переходит офис-менеджеру на этап физической подписи.")
    add_image(doc, "10-secretary-inbox.png", "Список договоров, ожидающих действий офис-менеджера.")
    add_image(doc, "11-secretary-task-card.png", "Карточка этапа подписи: печатный пакет, подписанный экземпляр и итоговое решение.", width=7.6)
    add_steps(doc, [
        "Откройте договор двойным кликом.",
        "Нажмите «Распечатать договор», чтобы сформировать печатный пакет.",
        "Передайте распечатанный экземпляр генеральному директору на подпись.",
        "После подписания приложите скан подписанного экземпляра.",
        "Выберите итоговое решение: «Подписан» или «Не согласован». Если договор не подписан, укажите комментарий.",
        "Нажмите кнопку завершения подписи.",
    ])

    add_section_break(doc)
    doc.add_heading("10. Реестр договоров", level=1)
    doc.add_paragraph("Реестр используется для поиска и просмотра договоров. В нём видны статусы, типы, контрагенты и ссылка на подписанный скан, если файл уже приложен.")
    add_image(doc, "12-contract-registry.png", "Реестр договоров: поиск, фильтры, статусы и колонка файла.")
    add_bullets(doc, [
        "Используйте поиск по номеру договора, контрагенту или ИНН.",
        "Фильтр «Статус» помогает найти договоры на согласовании, на подписании, подписанные или отклонённые.",
        "Дополнительные соглашения отображаются рядом с основным договором или раскрываются из строки основного договора.",
        "Если в колонке «Файл» есть ссылка «скан», её можно открыть для просмотра подписанного экземпляра.",
    ])

    doc.add_heading("11. Частые ситуации", level=1)
    add_bullets(doc, [
        "ФНС не вернула данные: заполните карточку контрагента вручную и продолжайте создание договора.",
        "Нужно добавить недостающий файл: используйте обсуждение или прикрепление файла в доступном блоке карточки.",
        "Поставлена виза с замечаниями или отрицательная виза: договор всё равно дойдёт до офис-менеджера после завершения обязательного круга.",
        "Договор уже подписан ранее: используйте импорт подписанного договора или импорт подписанного дополнительного соглашения.",
        "Письмо о задаче ведёт на конкретную карточку. Если вход не выполнен, после авторизации откроется нужный экран.",
    ])

    doc.add_heading("12. Короткая памятка", level=1)
    rows = [
        ("Создать договор", "«Мои договоры» → «Добавить» → «Договор»"),
        ("Создать дополнительное соглашение", "«Мои договоры» → «Добавить» → «Доп. соглашение»"),
        ("Импортировать подписанный документ", "«Мои договоры» → «Импорт» → нужный вариант импорта"),
        ("Открыть карточку", "Двойной клик по строке договора"),
        ("Поставить визу", "Открыть карточку → выбрать решение → при необходимости комментарий → «Сохранить решение»"),
        ("Написать в обсуждение", "Открыть карточку → блок «Обсуждение» → сообщение → «Отправить»"),
        ("Завершить подпись", "Офис-менеджер: распечатать пакет → получить подпись → приложить скан → выбрать итоговое решение"),
    ]
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [3600, 9840])
    for i, h in enumerate(["Задача", "Где выполнить"]):
        cell = table.cell(0, i)
        set_cell_shading(cell, LIGHT_BLUE)
        run = cell.paragraphs[0].add_run(h)
        set_run(run, size=9, color=DARK_BLUE, bold=True)
    for task, where in rows:
        cells = table.add_row().cells
        cells[0].text = task
        cells[1].text = where
        for cell in cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    set_run(r, size=9)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    DOCS_OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    doc.save(DOCS_OUTPUT)


if __name__ == "__main__":
    build()
